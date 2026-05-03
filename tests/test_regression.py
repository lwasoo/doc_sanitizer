from __future__ import annotations

import json
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document

from doc_sanitizer.document_io import apply_mapping_to_file, restore_file
from doc_sanitizer.fuzzy_mapping import (
    build_external_ai_prompt_sections,
    payload_from_json_text,
    suggest_placeholder_repairs,
)
from doc_sanitizer.mapping import mapping_entries, read_mapping
from gui_app.update_checker import (
    ReleaseAsset,
    UpdateInfo,
    compare_versions,
    download_release_asset,
    sanitize_filename,
    unique_path,
)


class MappingAndPromptTests(unittest.TestCase):
    def test_prompt_copy_section_excludes_sensitive_originals(self) -> None:
        # 通过用例：外部 AI 可复制区只能包含占位符，不能泄露原始敏感词。
        # 内部审核区可以保留原始名称线索，供用户自己判断归组是否合理。
        payload = payload_from_json_text(
            json.dumps(
                {
                    "entries": [
                        {
                            "placeholder": "__COMPANY_001__",
                            "original": "Acme",
                            "category": "COMPANY",
                            "enabled": True,
                        },
                        {
                            "placeholder": "__COMPANY_002__",
                            "original": "Acme Technology Co Ltd",
                            "category": "COMPANY",
                            "enabled": True,
                        },
                    ]
                },
                ensure_ascii=False,
            )
        )

        prompt, audit = build_external_ai_prompt_sections(payload)

        self.assertIn("__COMPANY_002__ / __COMPANY_001__", prompt)
        self.assertNotIn("Acme", prompt)
        self.assertNotIn("映射摘要", prompt)
        self.assertNotIn("人工确认", prompt)
        self.assertIn("Acme Technology Co Ltd", audit)

    def test_placeholder_repair_scores_split_auto_and_confirmation_cases(self) -> None:
        # 通过用例：轻微损坏的占位符应进入高置信自动修复；
        # 缺字/编号简写这类更可疑的写法只进入“需要用户确认”的分数段。
        payload = payload_from_json_text(
            json.dumps(
                {
                    "entries": [
                        {
                            "placeholder": "__COMPANY_001__",
                            "original": "Acme",
                            "category": "COMPANY",
                            "enabled": True,
                        }
                    ]
                },
                ensure_ascii=False,
            )
        )

        repairs = suggest_placeholder_repairs("Keep COMPANY_001 and maybe COMY_01.", mapping_entries(payload), min_score=0.70)
        scores = {repair.token: repair.score for repair in repairs}

        self.assertGreaterEqual(scores["COMPANY_001"], 0.90)
        self.assertGreaterEqual(scores["COMY_01"], 0.70)
        self.assertLess(scores["COMY_01"], 0.90)

    def test_invalid_mapping_json_is_rejected(self) -> None:
        # 失败用例：外部输入不是映射对象时应直接报错，避免生成误导性的 Prompt。
        with self.assertRaises(ValueError):
            payload_from_json_text(json.dumps(["not", "a", "mapping"]))

        # 失败用例：JSON 结构存在但没有任何有效 entries/replacements，也应报错。
        with self.assertRaises(ValueError):
            payload_from_json_text(json.dumps({"entries": []}))

    def test_very_weak_placeholder_match_is_not_suggested(self) -> None:
        # 失败用例：太不像的 token 不应被猜测成某个占位符，避免错误还原。
        payload = payload_from_json_text(
            json.dumps(
                {
                    "entries": [
                        {
                            "placeholder": "__COMPANY_001__",
                            "original": "Acme",
                            "category": "COMPANY",
                            "enabled": True,
                        }
                    ]
                },
                ensure_ascii=False,
            )
        )

        repairs = suggest_placeholder_repairs("This token ABC_99 should not match.", mapping_entries(payload), min_score=0.70)
        self.assertEqual(repairs, [])


class DocumentRoundTripTests(unittest.TestCase):
    def test_docx_sanitize_and_restore_with_confirmed_placeholder_repairs(self) -> None:
        # 通过用例：覆盖最核心的 docx 链路：
        # 原文 -> 脱敏文档/映射 JSON -> 外部 AI 改坏占位符 -> 按确认结果还原。
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source_path = root / "source.docx"
            sanitized_path = root / "sanitized.docx"
            edited_path = root / "edited.docx"
            restored_path = root / "restored.docx"
            mapping_path = root / "mapping.json"

            doc = Document()
            doc.add_paragraph("Acme Technology Co Ltd signed Project Phoenix.")
            doc.save(source_path)

            payload = {
                "version": 2,
                "entries": [
                    {
                        "placeholder": "__COMPANY_001__",
                        "original": "Acme Technology Co Ltd",
                        "category": "COMPANY",
                        "enabled": True,
                        "source": "test",
                    },
                    {
                        "placeholder": "__PROJECT_001__",
                        "original": "Project Phoenix",
                        "category": "PROJECT",
                        "enabled": True,
                        "source": "test",
                    },
                ],
            }

            apply_mapping_to_file(source_path, sanitized_path, payload, mapping_path)
            sanitized_text = "\n".join(p.text for p in Document(sanitized_path).paragraphs)
            self.assertIn("__COMPANY_001__", sanitized_text)
            self.assertIn("__PROJECT_001__", sanitized_text)
            self.assertEqual(read_mapping(mapping_path)["replacements"]["__COMPANY_001__"], "Acme Technology Co Ltd")

            edited = Document()
            edited.add_paragraph("COMPANY_001 signed __PROJECT-001__.")
            edited.save(edited_path)

            restore_file(
                input_path=edited_path,
                output_path=restored_path,
                mapping_path=mapping_path,
                placeholder_repairs={
                    "COMPANY_001": "__COMPANY_001__",
                    "__PROJECT-001__": "__PROJECT_001__",
                },
            )
            restored_text = "\n".join(p.text for p in Document(restored_path).paragraphs)
            self.assertIn("Acme Technology Co Ltd", restored_text)
            self.assertIn("Project Phoenix", restored_text)

    def test_restore_rejects_empty_mapping_file(self) -> None:
        # 失败用例：映射 JSON 没有有效条目时，不能生成一个看似成功但未还原的文件。
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            input_path = root / "input.docx"
            output_path = root / "output.docx"
            mapping_path = root / "mapping.json"

            doc = Document()
            doc.add_paragraph("__COMPANY_001__")
            doc.save(input_path)
            mapping_path.write_text(json.dumps({"entries": []}), encoding="utf-8")

            with self.assertRaises(ValueError):
                restore_file(input_path=input_path, output_path=output_path, mapping_path=mapping_path)


class UpdateCheckerTests(unittest.TestCase):
    def test_update_asset_selection_and_path_helpers(self) -> None:
        # 通过用例：同一个 Release 里有多个平台产物时，应按当前系统选择下载包。
        assets = [
            ReleaseAsset("v1.2.0-FileToolbox-macos.zip", "https://example.test/mac.zip"),
            ReleaseAsset("v1.2.0-FileToolbox.exe", "https://example.test/win.exe"),
        ]
        info = UpdateInfo(
            current_version="1.1.0",
            latest_version="1.2.0",
            release_name="v1.2.0",
            release_url="https://example.test/release",
            published_at="",
            assets=assets,
            is_update_available=True,
        )

        with patch.object(sys, "platform", "win32"):
            self.assertEqual(info.preferred_asset_name, "v1.2.0-FileToolbox.exe")
        with patch.object(sys, "platform", "darwin"):
            self.assertEqual(info.preferred_asset_name, "v1.2.0-FileToolbox-macos.zip")

        self.assertGreater(compare_versions("1.2.0", "1.1.9"), 0)
        self.assertEqual(sanitize_filename('bad:name*.exe'), "bad_name_.exe")
        with TemporaryDirectory() as temp_dir:
            existing = Path(temp_dir) / "FileToolbox.exe"
            existing.write_text("old", encoding="utf-8")
            self.assertEqual(unique_path(existing).name, "FileToolbox_2.exe")

    def test_update_download_rejects_release_without_assets(self) -> None:
        # 失败用例：Release 没有任何资产时不能假装下载成功，应提示用户打开 Release 页面人工查看。
        info = UpdateInfo(
            current_version="1.1.0",
            latest_version="1.2.0",
            release_name="v1.2.0",
            release_url="https://example.test/release",
            published_at="",
            assets=[],
            is_update_available=True,
        )

        with self.assertRaises(RuntimeError):
            download_release_asset(info)

    def test_unknown_version_string_compares_as_zero(self) -> None:
        # 失败/边界用例：无法解析数字的版本号按 0 处理，不应抛异常影响启动检测。
        self.assertEqual(compare_versions("dev", "0.0.0"), 0)
        self.assertLess(compare_versions("dev", "1.0.0"), 0)


if __name__ == "__main__":
    unittest.main()
